import os
import io
import json
import base64
import re  # <-- Importado para trabalharmos com Regex
import xml.etree.ElementTree as ET
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Any  # <-- Importado para dar flexibilidade ao payload
import pypdf
from docx import Document
from google import genai
from google.genai import types
from google.genai.errors import APIError

app = FastAPI(title="API Inteligente de Triagem e Extração de Notas/Faturas/Boletos - Gemini 2.5")

# Inicializa o cliente do Gemini
client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

# EXPANDIDO: Inclui termos comuns de Notas Fiscais, Faturas e agora de BOLETOS/COBRANÇAS
PALAVRAS_CHAVE_VALIDACAO = [
    "nota fiscal", "nf-e", "nfe", "danfe", "nfse", "prestador", "tomador", 
    "faturamento", "emissão", "valor recebido", "fatura", "invoice", 
    "duplicata", "comprovante de cobrança", "cobrança", "recibo",
    "boleto", "vencimento", "linha digitável", "código de barras", "cedente", 
    "beneficiário", "pagável", "sacado", "agência/código", "nosso número"
]

MESES_MAP = {
    "01": "Janeiro", "02": "Fevereiro", "03": "Março", "04": "Abril",
    "05": "Maio", "06": "Junho", "07": "Julho", "08": "Agosto",
    "09": "Setembro", "10": "Outubro", "11": "Novembro", "12": "Dezembro"
}

# --- AJUSTE NO PYDANTIC ---
class Payload(BaseModel):
    nome_arquivo: str
    pdf_base64: Any  # Aceita string pura ou objeto estruturado do Power Automate


# -------------------------------------------------------------------------
# FUNÇÕES DE FALLBACK COM REGEX (Plano de Contingência Local)
# -------------------------------------------------------------------------
def buscar_cnpj_fallback(texto: str) -> str:
    """Busca um CNPJ válido no texto que não seja o da Biochimico."""
    cnpj_biochimico = "33258401000448"
    padrao_cnpj = re.compile(r'\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}')
    encontrados = padrao_cnpj.findall(texto)
    
    for cnpj in encontrados:
        cnpj_limpo = re.sub(r'\D', '', cnpj)
        if cnpj_limpo != cnpj_biochimico:
            print(f"[FALLBACK REGEX] CNPJ do Fornecedor localizado via Regex: {cnpj}")
            return cnpj
            
    return ""

def buscar_valor_fallback(texto: str) -> str:
    """Busca o valor total baseado em linhas com palavras-chave fiscais ou bancárias."""
    linhas = texto.split('\n')
    padrao_valor = re.compile(r'\b\d{1,3}(?:\.\d{3})*,\d{2}\b')
    
    palavras_alvo = ["total", "valor liquido", "total da fatura", "valor total", "a pagar", "valor do documento", "(=) valor cobrado"]
    
    for linha in lines:
        linha_lower = linha.lower()
        if any(palavra in linha_lower for palavra in palavras_alvo):
            valores = padrao_valor.findall(linha)
            if valores:
                print(f"[FALLBACK REGEX] Valor localizado via Regex na linha [{linha.strip()}]: {valores[-1]}")
                return valores[-1]
                
    return ""


@app.get("/")
def health_check():
    chave_existe = "SIM" if os.environ.get("GEMINI_API_KEY") else "NÃO"
    return {
        "status": "online", 
        "servico": "API Triagem e Extração Avançada Ativa (Suporta Boletos)",
        "chave_configurada": chave_existe
    }

@app.post("/extrair-nf")
def extrair_nf(payload: Payload):
    # CORREÇÃO CRÍTICA: Trata espaços em branco nas pontas e joga toda a string para minúsculo
    nome_arq = payload.nome_arquivo.strip().lower()
    print(f"=== [API] Nova requisição: Arquivo [{payload.nome_arquivo}] ===")
    
    # -------------------------------------------------------------------------
    # FILTRO 1: VALIDAÇÃO DE EXTENSÃO (Otimizada com Expressão Regular)
    # -------------------------------------------------------------------------
    # O padrão r'\.(pdf|xml|docx)$' valida de forma resiliente o final real do arquivo
    if not re.search(r'\.(pdf|xml|docx)$', nome_arq):
        print(f"[TRIAGEM] Arquivo descartado por extensão inválida após tratamento: {nome_arq}")
        raise HTTPException(
            status_code=422, 
            detail=f"Arquivo descartado: Extensão não permitida. Recebido: '{payload.nome_arquivo}'. Use apenas PDF, XML ou DOCX."
        )

    # -------------------------------------------------------------------------
    # DECODIFICAÇÃO DO BASE64 COM TRATAMENTO PARA POWER AUTOMATE
    # -------------------------------------------------------------------------
    try:
        raw_base64 = payload.pdf_base64
        if isinstance(raw_base64, dict):
            if "$content" in raw_base64:
                raw_base64 = raw_base64["$content"]
            elif "contentBytes" in raw_base64:
                raw_base64 = raw_base64["contentBytes"]
        
        string_base64_limpa = str(raw_base64).strip().replace("\n", "").replace("\r", "").strip()
        bytes_ascii = string_base64_limpa.encode('ascii', errors='ignore')
        conteudo_bytes = base64.b64decode(bytes_ascii)
        
    except Exception as e:
        print(f"[ERRO CRÍTICO BASE64] String inválida recebida: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Erro ao decodificar Base64: {str(e)}")

    # -------------------------------------------------------------------------
    # FLUXO AUTOMÁTICO: PROCESSAMENTO DE XML (Custo Zero de IA)
    # -------------------------------------------------------------------------
    if nome_arq.endswith(".xml"):
        print("[PROCESSAMENTO] Identificado arquivo XML. Processando nativamente...")
        try:
            string_xml = conteudo_bytes.decode('utf-8', errors='ignore').strip()
            raiz = ET.fromstring(string_xml)
            
            for elem in raiz.iter():
                if '}' in elem.tag:
                    elem.tag = elem.tag.split('}', 1)[1]
            
            if raiz.find('.//infNfe') is None and raiz.find('.//infNFSe') is None and 'nfe' not in raiz.tag.lower():
                print(f"[TRIAGEM] XML rejeitado. Não possui estrutura fiscal: {payload.nome_arquivo}")
                raise HTTPException(status_code=422, detail="Arquivo descartado: O XML enviado não é um documento fiscal válido.")

            numero_nf = raiz.find('.//ide/nNF') or raiz.find('.//numero')
            emitente = raiz.find('.//emit/xNome') or raiz.find('.//prestadorServico/razaoSocial') or raiz.find('.//emit/xFant')
            cnpj = raiz.find('.//emit/CNPJ') or raiz.find('.//emit/CPF') or raiz.find('.//prestadorServico/identificacaoPrestador/cnpj')
            valor = raiz.find('.//vNF') or raiz.find('.//valores/valorLiquido') or raiz.find('.//vProd')
            data = raiz.find('.//dhEmi') or raiz.find('.//dataEmissao') or raiz.find('.//dEmi')

            mes_extenso = "Mês Não Identificado"
            if data is not None and data.text:
                if '-' in data.text:
                    partes_data = data.text.split('-')
                    if len(partes_data) >= 2 and partes_data[1] in MESES_MAP:
                        mes_extenso = MESES_MAP[partes_data[1]]
                elif '/' in data.text:
                    partes_barra = data.text.split('/')
                    if len(partes_barra) >= 2 and partes_barra[1] in MESES_MAP:
                        mes_extenso = MESES_MAP[partes_barra[1]]

            resultado_xml = {
                "tipo_documento": "Nota Fiscal",
                "fornecedor": emitente.text if emitente is not None else "Não encontrado",
                "cnpj_cpf_nif": cnpj.text if cnpj is not None else "Não encontrado",
                "numero_nf": numero_nf.text if numero_nf is not None else "Não encontrado",
                "data_emissao": data.text if data is not None else "Não encontrado",
                "mes_extenso": mes_extenso,
                "valor_total": valor.text if valor is not None else "0.00"
            }
            print(f"[SUCESSO LOCAL] XML extraído com sucesso: {resultado_xml}")
            return resultado_xml

        except HTTPException as http_err:
            raise http_err
        except Exception as e:
            print(f"[ERRO XML] Falha no parse do XML: {str(e)}")
            raise HTTPException(status_code=422, detail=f"Falha ao ler a estrutura do arquivo XML: {str(e)}")

    # -------------------------------------------------------------------------
    # FLUXO TEXTUAL: EXTRAÇÃO DE TEXTO PARA DOCX OU PDF (COM SUPORTE A TABELAS)
    # -------------------------------------------------------------------------
    texto_extraido = ""
    
    if ".docx" in nome_arq:
        print("[PROCESSAMENTO] Extraindo texto de documento Word (.docx)...")
        try:
            doc = Document(io.BytesIO(conteudo_bytes))
            for paragrafo in doc.paragraphs:
                if paragrafo.text.strip():
                    texto_extraido += paragrafo.text + "\n"
            
            for tabela in doc.tables:
                for linha in tabela.rows:
                    texto_linha = [celula.text.strip() for celula in linha.cells if celula.text.strip()]
                    if texto_linha:
                        texto_extraido += " | ".join(list(dict.fromkeys(texto_linha))) + "\n"
                        
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Erro ao ler arquivo Word: {str(e)}")

    elif ".pdf" in nome_arq:
        print("[PROCESSAMENTO] Extraindo texto de arquivo PDF...")
        try:
            leitor = pypdf.PdfReader(io.BytesIO(conteudo_bytes))
            for pagina in leitor.pages:
                texto_extraido += pagina.extract_text() or ""
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Erro ao ler arquivo PDF: {str(e)}")

    # -------------------------------------------------------------------------
    # FILTRO 2: VALIDAÇÃO TEXTUAL (Mínimo de 1 termo fiscal/bancário presente)
    # -------------------------------------------------------------------------
    texto_analise = texto_extraido.lower()
    matches = [termo for termo in PALAVRAS_CHAVE_VALIDACAO if termo in texto_analise]
    
    if len(matches) < 1:
        print(f"[TRIAGEM] Arquivo textual rejeitado por falta de contexto fiscal/bancário ({len(matches)} termos encontrados).")
        raise HTTPException(
            status_code=422, 
            detail="Arquivo descartado: O conteúdo não condiz com uma Nota Fiscal, Fatura ou Boleto legítimo."
        )

    # -------------------------------------------------------------------------
    # CHAMADA GEMINI
    # -------------------------------------------------------------------------
    try:
        print("[GEMINI] Documento validado! Enviando texto para estruturação avançada...")
        
        prompt = (
            "Atue como um analista fiscal e financeiro especialista em extração de dados de documentos brasileiros.\n"
            "Analise textualmente o documento fornecido. Classifique o tipo de documento entre 'Nota Fiscal', 'Fatura' ou 'Boleto' "
            "e extraia as informações necessárias estruturadas estritamente no formato JSON solicitado.\n\n"
            
            "DIRETRIZES CRÍTICAS PARA GARANTIR A CAPTURA:\n"
            "1. **tipo_documento**: Defina rigidamente se o documento é uma 'Nota Fiscal', 'Fatura' ou 'Boleto'.\n\n"
            
            "2. **fornecedor**: Identifique a empresa EMISSORA, PRESTADORA ou BENEFICIÁRIA/CEDENTE da cobrança.\n"
            "   - REGRA DE EXCLUSÃO MANDATÓRIA: A empresa 'BIOCHIMICO' (or 'INSTITUTO BIOCHIMICO') é estritamente o CLIENTE/TOMADOR/PAGADOR. Portanto, NUNCA capture 'BIOCHIMICO' neste campo.\n"
            "   - O fornecedor real será a OUTRA empresa que aparece no texto (geralmente associada a termos como 'Emitente', 'Prestador', 'Beneficiário', 'Cedente' ou listada no topo do documento).\n"
            "   - Caso venha em formato de tabela misturado, separe e traga apenas o nome da empresa parceira.\n\n"
            
            "3. **cnpj_cpf_nif**: Extraia o CNPJ pertencente ao fornecedor/beneficiário identificado no item anterior.\n"
            "   - REGRA DE EXCLUSÃO MANDATÓRIA: O CNPJ '33.258.401/0004-48' pertence à Biochimico. NUNCA retorne este número aqui.\n"
            "   - Localize o outro CNPJ ou CPF presente no documento correspondente à empresa emissora da cobrança e retorne-o formatado.\n\n"
            
            "4. **numero_nf**: Encontre o número do documento fiscal, número da faturamento ou o 'Número do Documento' / 'Nosso Número' se for um boleto. "
            "Se o número constar no nome do arquivo (ex: '602261'), certifique-se de validar se ele bate com o sequencial encontrado no texto.\n\n"
            
            "5. **data_emissao**: Capture a data de emissão ou data do documento expressa no texto. Caso o documento seja um boleto e não declare explicitamente a data de emissão, use a data de vencimento disponível. "
            "Garante o formato de ano com 4 dígitos (Ex: '29/05/2026').\n\n"
            
            "6. **mes_extenso**: Baseado exclusivamente na data tratada no item anterior, determine o mês por extenso em português, com a inicial maiúscula (Ex: 'Maio').\n\n"
            
            "7. **valor_total**: Localize o valor monetário final do documento (palavras-chave como 'TOTAL', 'VALOR', 'VALOR DO DOCUMENTO' ou 'COBRADO'). "
            "Retorne apenas os caracteres numéricos e a vírgula/ponto decimal (Ex: '9800,50').\n\n"
            
            f"Texto do documento para análise:\n{texto_extraido}"
        )

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=types.Schema(
                    type=types.Type.OBJECT,
                    properties={
                        "tipo_documento": types.Schema(type=types.Type.STRING),
                        "fornecedor": types.Schema(type=types.Type.STRING),
                        "cnpj_cpf_nif": types.Schema(type=types.Type.STRING),
                        "numero_nf": types.Schema(type=types.Type.STRING),
                        "data_emissao": types.Schema(type=types.Type.STRING),
                        "mes_extenso": types.Schema(type=types.Type.STRING),
                        "valor_total": types.Schema(type=types.Type.STRING),
                    },
                    required=["tipo_documento", "fornecedor", "cnpj_cpf_nif", "numero_nf", "data_emissao", "mes_extenso", "valor_total"],
                ),
            ),
        )
        
        texto_resposta = response.text.strip() if response.text else "RESPOSTA_VAZIA"
        resultado_ia = json.loads(texto_resposta)

        # -------------------------------------------------------------------------
        # CAMADA PÓS-PROCESSAMENTO: APLICAÇÃO DOS FALLBACKS
        # -------------------------------------------------------------------------
        if not resultado_ia.get("cnpj_cpf_nif") or resultado_ia.get("cnpj_cpf_nif").strip() in ["", "Não encontrado", "Não Informado"]:
            cnpj_encontrado = buscar_cnpj_fallback(texto_extraido)
            if cnpj_encontrado:
                resultado_ia["cnpj_cpf_nif"] = cnpj_encontrado

        if not resultado_ia.get("valor_total") or resultado_ia.get("valor_total").strip() in ["", "0.00", "0", "Não encontrado"]:
            valor_encontrado = buscar_valor_fallback(texto_extraido)
            if valor_encontrado:
                resultado_ia["valor_total"] = valor_encontrado

        if not resultado_ia.get("fornecedor") or resultado_ia.get("fornecedor").strip() == "":
            resultado_ia["fornecedor"] = "Fornecedor/Beneficiário Não Identificado"

        print(f"[GEMINI + FALLBACK] Mapeamento bem-sucedido: {resultado_ia.get('fornecedor')} | Tipo: {resultado_ia.get('tipo_documento')} | Valor: {resultado_ia.get('valor_total')}")
        return resultado_ia

    except APIError as api_err:
        if api_err.code == 429:
            print("[ALERTA COTA] Limite atingido no Gemini.")
            raise HTTPException(status_code=429, detail="Limite de requisições atingido. Tente novamente em breve.")
        raise HTTPException(status_code=500, detail=f"Erro na API do Google: {str(api_err)}")
    except json.JSONDecodeError as json_err:
        print(f"[ERRO JSON PARSE] Resposta bruta da IA:\n{texto_resposta}")
        raise HTTPException(status_code=500, detail="Erro interno: A IA não retornou um formato estruturado válido.")
    except Exception as e:
        print(f"[ERRO GERAL] Falha: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erro no processamento da IA: {str(e)}")
